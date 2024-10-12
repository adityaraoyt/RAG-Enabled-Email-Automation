from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.core import Settings, SimpleDirectoryReader, VectorStoreIndex, ServiceContext
from llama_index.core.retrievers import VectorIndexRetriever
from llama_index.core.query_engine import RetrieverQueryEngine
from llama_index.core.postprocessor import SimilarityPostprocessor
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.core import Settings, VectorStoreIndex,SimpleDirectoryReader,ServiceContext,PromptTemplate
from llama_index.llms.huggingface import HuggingFaceLLM
from transformers import AutoModelForCausalLM, AutoTokenizer
from docx import Document
import os
import re
import argparse
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from getpass import getpass
from datetime import datetime

# Set up the LLM
model_name = "TinyLlama/TinyLlama-1.1B-Chat-v1.0"
model = AutoModelForCausalLM.from_pretrained(model_name,
                                             device_map="auto",
                                             trust_remote_code=False,
                                             revision="main")

tokenizer = AutoTokenizer.from_pretrained(model_name, use_fast=True)

# Set up the embedding model
embed_model = HuggingFaceEmbedding(model_name="BAAI/bge-small-en-v1.5")
Settings.chunk_size = 256
Settings.chunk_overlap = 25

# Set up HuggingFaceLLM
llm = HuggingFaceLLM(
    context_window=1024,
    max_new_tokens=500,  
    generate_kwargs={"temperature": 0.2, "do_sample": False},
    tokenizer=tokenizer,
    model=model
)

def load_company_documents(company_name):
    company_dir = os.path.join('docs', company_name.lower())
    if not os.path.exists(company_dir):
        print(f"Error: Directory for {company_name} not found.")
        return None

    try:
        documents = SimpleDirectoryReader(company_dir).load_data()
        print(f"Loaded {len(documents)} documents for {company_name}")
        return documents
    except Exception as e:
        print(f"Error loading documents for {company_name}: {e}")
        return None

def read_word_file(file_path):
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

MyCompany_services_path = 'MyCompany/MyCompany.docx'
MyCompany_services = read_word_file(MyCompany_services_path)

def generate_email(days, company_name, recipient_email, index):
    # Query the index for relevant information about the company
    query = f"How can MyCompany's services benefit {company_name}?"
    retriever = index.as_retriever(similarity_top_k=1)
    retrieved_nodes = retriever.retrieve(query)
    context = "\n".join([node.text for node in retrieved_nodes])

    # Determine max word count based on days
    if days <= 7:
        max_words = 400
    elif days <= 14:
        max_words = 300
    elif days <= 21:
        max_words = 200
    elif days <= 29:
        max_words = 100
    else:
        max_words = 50

    prompt_template = f"""

Services provided by MyCompany: {MyCompany_services[:200]}

{company_name} Context: {context[:200]}

Write a {max_words}-word follow-up email from MyCompany to {recipient_email} at {company_name} after {days} days of no response. Talk about 
how {company_name} would benefit from using MyCompany's services and also suggest next steps that MyCompany plans on taking with company1 for their partnership.

Be professional, friendly, and specific. No placeholders. Keep the email under { max_words } words long.

Email:
"""
    inputs = tokenizer(prompt_template, return_tensors="pt")
    outputs = model.generate(
        input_ids=inputs["input_ids"].to("cuda"),
        max_new_tokens=1000,
        do_sample=False,
        repetition_penalty=1.2,
        no_repeat_ngram_size=3
    )

    response = tokenizer.decode(outputs[0], skip_special_tokens=True)
    
    return response

def format_email(text):
    # Extract the generated email content
    email_content = text.split("Generate the email now:")[-1].strip()
    
    # Separate subject line
    subject_match = re.search(r"Subject:(.+)", email_content, re.IGNORECASE)
    subject = subject_match.group(1).strip() if subject_match else "Follow-up"
    
    # Remove subject line from content
    email_body = re.sub(r"Subject:.+\n", "", email_content, flags=re.IGNORECASE)
    
    # Ensure proper line breaks
    email_body = re.sub(r"\n{3,}", "\n\n", email_body)
    
    # Format the email
    formatted_email = f"Subject: {subject}\n\n{email_body}"
    
    return subject, email_body

def save_email_to_docx(subject, body, sender_email, recipient_email, company_name):
    doc = Document()
    doc.add_heading('Email Content', 0)
    
    doc.add_paragraph(f'From: {sender_email}')
    doc.add_paragraph(f'To: {recipient_email}')
    doc.add_paragraph(f'Company: {company_name}')
    doc.add_paragraph(f'Subject: {subject}')
    doc.add_paragraph('\n')
    doc.add_paragraph(body)
    
    # Create 'responses' folder if it doesn't exist
    if not os.path.exists('responses'):
        os.makedirs('responses')
    
    # Generate a unique filename
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f'responses/email_response_{company_name}_{timestamp}.docx'
    
    doc.save(filename)
    print(f"Email content saved to {filename}")

def send_email(sender_email, sender_password, recipient_email, subject, body):
    msg = MIMEMultipart()
    msg['From'] = sender_email
    msg['To'] = recipient_email
    msg['Subject'] = subject

    msg.attach(MIMEText(body, 'plain'))

    try:
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(sender_email, sender_password)
        text = msg.as_string()
        server.sendmail(sender_email, recipient_email, text)
        server.quit()
        print("Email sent successfully!")
    except Exception as e:
        print(f"An error occurred while sending the email: {str(e)}")

def validate_email(email, company_name):
    # Convert to lowercase for case-insensitive comparison
    email = email.lower()
    company_name = company_name.lower()
    
    # Remove spaces and special characters from company name
    company_name = re.sub(r'[^a-zA-Z0-9]', '', company_name)
    
    # Create the pattern
    pattern = rf'^[a-zA-Z0-9._%+-]+@{company_name}\.[a-zA-Z]+$'
    
    if re.match(pattern, email):
        return True
    else:
        return False

import argparse
from getpass import getpass

def main():
    parser = argparse.ArgumentParser(description="Generate and send a follow-up email.")
    parser.add_argument("--days", type=int, required=True, help="Number of days since last contact")
    parser.add_argument("--company", type=str, required=True, help="Target company name")
    parser.add_argument("--email", type=str, required=True, help="Recipient email address")
    parser.add_argument("--sender-email", type=str, required=True, help="Sender's email address")
    parser.add_argument("--send", action="store_true", help="Send the email immediately")
    args = parser.parse_args()

    
    company_name = args.company
    days = args.days
    recipient_email = args.email
    sender_email = args.sender_email

    # Validate email
    if not validate_email(recipient_email, company_name):
        print(f"Invalid email format. The email should end with '@{company_name}.<something>'")
        return

    # Load company-specific documents
    documents = load_company_documents(company_name)
    if not documents:
        print("Exiting due to document loading error.")
        return

    # Create the index
    index = VectorStoreIndex.from_documents(documents, embed_model=embed_model)

    email_content = generate_email(days, company_name, recipient_email, index)
    subject, body = format_email(email_content)
    
    print("\nGenerated Email:\n")
    print(f"Subject: {subject}\n")
    print(body)

    # Save email content to docx file
    save_email_to_docx(subject, body, sender_email, recipient_email, company_name)

    if args.send:
        sender_password = getpass("Enter your email password: ")
        send_email(sender_email, sender_password, recipient_email, subject, body)
    else:
        print("\nEmail saved to docx file. Use --send flag to send the email.")

if __name__ == "__main__":
    main()

